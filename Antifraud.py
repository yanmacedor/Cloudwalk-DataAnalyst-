import os
import pandas as pd
import numpy as np
from datetime import datetime
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from sklearn.preprocessing import StandardScaler
from sklearn.cluster import KMeans, DBSCAN, kmeans_plusplus
from sklearn.metrics import silhouette_score
from openpyxl.drawing.image import Image as XLImage
from openpyxl.utils import get_column_letter
from docx import Document
import seaborn as sns

csv_file = r"C:\Users\Administrator\OneDrive\Área de Trabalho\a\transactional-sample.csv"
base_output_dir = r"C:\Users\Administrator\OneDrive\Área de Trabalho\a"
timestamp = datetime.now().strftime('%Y%m%d_%H%M')
output_dir = os.path.join(base_output_dir, timestamp)
os.makedirs(output_dir, exist_ok=True)
log_file = os.path.join(output_dir, "log_transacoes.xlsx")

def build_user_profiles(df: pd.DataFrame) -> dict:
    user_stats = df.groupby('user_id').agg(
        avg_amount=('transaction_amount', 'mean'),
        std_dev_amount=('transaction_amount', 'std'),
        total_transactions=('transaction_id', 'count'),
        device_count=('device_id', 'nunique')
    )
    user_stats['std_dev_amount'] = user_stats['std_dev_amount'].fillna(0)
    percentile_95 = df.groupby('user_id')['transaction_amount'].apply(lambda x: np.percentile(x, 95))
    min_interval = (
        df.sort_values('transaction_date')
        .groupby('user_id')['transaction_date']
        .diff()
        .dt.total_seconds()
        .groupby(df['user_id'])
        .min()
        .fillna(0)
    )
    profiles = {}
    for user_id in user_stats.index:
        profiles[user_id] = {
            'avg_amount': user_stats.loc[user_id, 'avg_amount'],
            'std_dev_amount': user_stats.loc[user_id, 'std_dev_amount'],
            'total_transactions': user_stats.loc[user_id, 'total_transactions'],
            'device_count': user_stats.loc[user_id, 'device_count'],
            'percentile_95': percentile_95.get(user_id, user_stats.loc[user_id, 'avg_amount']),
            'min_interval': min_interval.get(user_id, None),
            'seen_devices': set(df[df['user_id'] == user_id]['device_id'].dropna())
        }
    return profiles

class AntiFraudSystem:
    def __init__(self, user_profiles: dict):
        self.user_profiles = user_profiles
        self.rules = {
            'std_dev_multiplier': 3.5,
            'max_transactions_per_hour': 5,
            'max_amount_per_day': 5000,
            'max_amount_single': 3000,
            'card_testing_amount': 5.00,
            'suspicious_intervals': {'critical': 2, 'high': 10},
            'blocked_devices': set(),
            'chargeback_users': set(),
            'user_hour_count': {},
            'user_day_amount': {},
            'user_last_tx': {}
        }
        self.decision_log = []

    def check_transaction(self, transaction: dict) -> dict:
        user_id = transaction['user_id']
        amount = transaction['transaction_amount']
        device_id = transaction.get('device_id')
        now = datetime.now()
        risk_score, reasons = 0, []
        if user_id in self.rules['chargeback_users']:
            return {"recommendation": "deny", "risk_score": 100, "reasons": ["Usuário em blacklist por chargeback anterior"]}
        if device_id in self.rules['blocked_devices']:
            return {"recommendation": "deny", "risk_score": 100, "reasons": ["Dispositivo com múltiplos chargebacks"]}
        profile = self.user_profiles.get(user_id)
        if profile:
            threshold = profile['percentile_95']
            if amount > threshold:
                risk_score += 45
                reasons.append(f"Valor R${amount:.2f} acima do percentil 95 (R${threshold:.2f})")
            if device_id and device_id not in profile['seen_devices']:
                risk_score += 40
                reasons.append("Novo dispositivo (risco de ATO)")
            if profile['device_count'] > 3:
                risk_score += 20
                reasons.append("Muitos dispositivos diferentes")
        hour_key = f"{user_id}_{now.strftime('%Y%m%d%H')}"
        self.rules['user_hour_count'][hour_key] = self.rules['user_hour_count'].get(hour_key, 0) + 1
        if self.rules['user_hour_count'][hour_key] > self.rules['max_transactions_per_hour']:
            risk_score += 30
            reasons.append("Excedeu transações/hora")
        day_key = f"{user_id}_{now.strftime('%Y%m%d')}"
        self.rules['user_day_amount'][day_key] = self.rules['user_day_amount'].get(day_key, 0) + amount
        if self.rules['user_day_amount'][day_key] > self.rules['max_amount_per_day']:
            risk_score += 35
            reasons.append("Excedeu limite diário")
        if amount > self.rules['max_amount_single']:
            risk_score += 50
            reasons.append("Valor único muito alto")
        last_tx_time = self.rules['user_last_tx'].get(user_id)
        self.rules['user_last_tx'][user_id] = now
        if last_tx_time:
            interval = (now - last_tx_time).total_seconds() / 60
            if interval < self.rules['suspicious_intervals']['critical']:
                risk_score += 40
                reasons.append("Intervalo criticamente curto")
            elif interval < self.rules['suspicious_intervals']['high']:
                risk_score += 20
                reasons.append("Intervalo suspeito")
        if amount < self.rules['card_testing_amount']:
            risk_score += 25
            reasons.append("Valor muito baixo (possível card testing)")
        final_score = min(risk_score, 100)
        if final_score >= 80:
            recommendation = "deny"
        elif final_score >= 40:
            recommendation = "review"
        else:
            recommendation = "approve"
        self.decision_log.append({
            "transaction_id": transaction['transaction_id'],
            "recommendation": recommendation,
            "risk_score": final_score,
            "reasons": reasons
        })
        return {"recommendation": recommendation, "risk_score": final_score, "reasons": reasons or ["Baixo risco"]}

def create_features_for_clustering(df: pd.DataFrame, user_profiles: dict):
    df['user_cbk_count'] = df.groupby('user_id')['has_cbk'].transform('sum')
    cashback_df = df[df['has_cbk'] == True].sort_values(['user_id', 'transaction_date'])
    cashback_df['interval'] = cashback_df.groupby('user_id')['transaction_date'].diff().dt.total_seconds()
    avg_interval = cashback_df.groupby('user_id')['interval'].mean().round(2)
    df['user_total_transactions'] = df['user_id'].map(lambda x: user_profiles.get(x, {}).get('total_transactions', 0))
    df['avg_time_between_cbk_min'] = df['user_id'].map(avg_interval)
    return df[['transaction_amount', 'user_total_transactions', 'user_cbk_count', 'avg_time_between_cbk_min']].fillna(0)

def run_clustering(df: pd.DataFrame, user_profiles: dict, writer: pd.ExcelWriter):
    features = create_features_for_clustering(df, user_profiles)
    scaler = StandardScaler()
    scaled_features = scaler.fit_transform(features)
    init_centers, _ = kmeans_plusplus(scaled_features, n_clusters=4, random_state=42)
    kmeans = KMeans(n_clusters=4, init=init_centers, n_init=1, random_state=42)
    df['cluster'] = kmeans.fit_predict(scaled_features)
    silhouette = silhouette_score(scaled_features, df['cluster'])
    blacklist_users = set(df[df['has_cbk'] == True]['user_id'])
    device_fraud_count = df[df['has_cbk'] == True].groupby('device_id').size()
    blacklist_devices = set(device_fraud_count[device_fraud_count > 2].index.dropna())
    df['blacklist'] = df.apply(
        lambda row: 'Usuário' if row['user_id'] in blacklist_users
        else 'Dispositivo' if row['device_id'] in blacklist_devices else '',
        axis=1
    )
    df['avg_cbk_per_transaction'] = (df['user_cbk_count'] / df['user_total_transactions']).replace([np.inf, -np.inf], 0)
    df['user_transaction_count'] = df.groupby('user_id')['transaction_id'].transform('count')
    df['first_transaction_date'] = df.groupby('user_id')['transaction_date'].transform('min')
    return df, kmeans

def analyze_transactions(df: pd.DataFrame, antifraud: AntiFraudSystem, writer: pd.ExcelWriter):
    results = []
    for _, row in df.iterrows():
        transaction = row.to_dict()
        result = antifraud.check_transaction(transaction)
        result["transaction_id"] = transaction["transaction_id"]
        results.append(result)
    log_df = pd.DataFrame(results)
    merged_df = df.merge(log_df, on='transaction_id')
    merged_df.to_excel(writer, sheet_name="Log de Transações", index=False)
    return merged_df

def save_analysis_to_docx(df: pd.DataFrame, filename=None):
    if filename is None:
        filename = os.path.join(output_dir, "analise_final.docx")
    else:
        if not os.path.isabs(filename):
            filename = os.path.join(output_dir, filename)
    suspeitos = [
        "1. Usuários com muitos dispositivos diferentes podem indicar fraude.",
        "2. Muitos chargebacks associados ao mesmo dispositivo ou usuário.",
        "3. Transações de valor muito alto ou muito baixo podem ser suspeitas.",
        "4. Intervalos curtos entre transações podem indicar automação ou teste de cartões."
    ]
    outros_dados = [
        "1. Localização geográfica do usuário/dispositivo (país, cidade, IP).",
        "2. Histórico de uso do dispositivo (outros usuários, frequência).",
        "3. Horário da transação (transações fora do horário comercial).",
        "4. Tipo de produto/serviço comprado.",
        "5. Dados do cartão (BIN, país de emissão, etc).",
        "6. Informações sobre o app ou versão do sistema operacional.",
        "7. Dados de login/autenticação (mudanças de senha, tentativas falhas)."
    ]
    resumo = [
        ("Total de transações", len(df)),
        ("Total de usuários", df['user_id'].nunique()),
        ("Total de dispositivos", df['device_id'].nunique()),
        ("Transações com chargeback", df['has_cbk'].sum()),
        ("Usuários com chargeback", df[df['has_cbk'] == True]['user_id'].nunique()),
        ("Maior valor de transação", df['transaction_amount'].max()),
        ("Menor valor de transação", df['transaction_amount'].min())
    ]
    doc = Document()
    doc.add_heading("Análise de Dados e Sugestões", level=1)
    doc.add_heading("Resumo dos Dados", level=2)
    for item, valor in resumo:
        doc.add_paragraph(f"{item}: {valor}")
    doc.add_heading("Padrões Suspeitos", level=2)
    for item in suspeitos:
        doc.add_paragraph(item, style='List Number')
    doc.add_heading("Outros Dados Úteis", level=2)
    for item in outros_dados:
        doc.add_paragraph(item, style='List Number')
    doc.save(filename)

def confusion_matrix_analysis(df: pd.DataFrame, writer: pd.ExcelWriter = None):
    antifraud_results = df[['has_cbk', 'recommendation']]
    tp = ((antifraud_results['has_cbk'] == True) & (antifraud_results['recommendation'] == 'deny')).sum()
    fn = ((antifraud_results['has_cbk'] == True) & (antifraud_results['recommendation'] == 'approve')).sum()
    fp = ((antifraud_results['has_cbk'] == False) & (antifraud_results['recommendation'] == 'deny')).sum()
    tn = ((antifraud_results['has_cbk'] == False) & (antifraud_results['recommendation'] == 'approve')).sum()
    total = tp + fn + fp + tn
    taxa_fp = fp / total if total > 0 else 0
    taxa_fn = fn / total if total > 0 else 0
    conf_matrix = [[tp, fp], [fn, tn]]
    plt.figure(figsize=(5, 4))
    sns.heatmap(conf_matrix, annot=True, fmt="d", cmap="Blues",
                xticklabels=["Fraude", "Não Fraude"],
                yticklabels=["Negado", "Aprovado"])
    plt.title("Matriz de Confusão")
    plt.xlabel("Realidade")
    plt.ylabel("Decisão do Sistema")
    plt.tight_layout()
    img_path = os.path.join(output_dir, "matriz_confusao.png")
    plt.savefig(img_path)
    plt.close()

def main():
    try:
        df = pd.read_csv(csv_file)
        df['transaction_date'] = pd.to_datetime(df['transaction_date'], errors='coerce')
        if 'has_cbk' not in df.columns:
            df['has_cbk'] = False
        df.fillna({'device_id': 'unknown', 'transaction_amount': 0, 'has_cbk': False}, inplace=True)
        user_profiles = build_user_profiles(df)
        antifraud = AntiFraudSystem(user_profiles)
        with pd.ExcelWriter(log_file, engine='openpyxl') as writer:
            df_clustered, kmeans = run_clustering(df, user_profiles, writer)
            merged_df = analyze_transactions(df_clustered, antifraud, writer)
            confusion_matrix_analysis(merged_df, writer)
        save_analysis_to_docx(df_clustered)
        print(f"\n--- ANÁLISE CONCLUÍDA ---")
    except FileNotFoundError:
        print(f"ERRO: O arquivo {csv_file} não foi encontrado.")
    except Exception as e:
        print(f"Ocorreu um erro inesperado: {e}")

if __name__ == "__main__":
    main()